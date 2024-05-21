from PyQt5.QtWidgets import QWidget, QVBoxLayout, QLabel, QLineEdit, QPushButton, QTextEdit, QFileDialog, QMessageBox
from docx import Document

class DocumentGenerator(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Generate Document")
        self.setGeometry(100, 100, 600, 400)
        self.initUI()

    def initUI(self):
        layout = QVBoxLayout()

        self.titleLabel = QLabel("Title:")
        self.titleInput = QLineEdit()
        layout.addWidget(self.titleLabel)
        layout.addWidget(self.titleInput)

        self.contentLabel = QLabel("Content:")
        self.contentInput = QTextEdit()
        layout.addWidget(self.contentLabel)
        layout.addWidget(self.contentInput)

        self.generateButton = QPushButton("Generate Document")
        self.generateButton.clicked.connect(self.generate_document)
        layout.addWidget(self.generateButton)

        self.setLayout(layout)

    def generate_document(self):
        title = self.titleInput.text()
        content = self.contentInput.toPlainText()

        if title and content:
            document = Document()
            document.add_heading(title, 0)
            document.add_paragraph(content)

            filePath, _ = QFileDialog.getSaveFileName(self, "Save Document", "", "Word Documents (*.docx)")
            if filePath:
                document.save(filePath)
                QMessageBox.information(self, "Success", "Document generated and saved successfully")
            else:
                QMessageBox.warning(self, "Error", "No file selected")
        else:
            QMessageBox.warning(self, "Error", "Title and content are required")

